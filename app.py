from flask import Flask, request, jsonify, send_file, render_template
import pdfplumber
from docx import Document
from flask_cors import CORS
import os
import openai

app = Flask(__name__)
openai.api_key = "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"

# ---------- New Home Route ----------

@app.route('/')
def home():
    return render_template('index.html')

# ---------- Utility Functions ----------

def extract_text(pdf_file):
    """
    Extract text content from a PDF file.
    """
    text = ""
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def generate_feedback(resume_text, job_role):
    """
    Generate basic feedback and enhanced resume content.
    """
    tips = (
        f"🔍 Tips for your {job_role} job search:\n"
        f"1️⃣ Use specific keywords from {job_role} job postings.\n"
        f"2️⃣ Highlight measurable achievements.\n"
        f"3️⃣ Make sure your LinkedIn matches your resume.\n"
    )
    new_resume_text = (
        f"Enhanced Resume Tailored for {job_role}:\n\n"
        f"{resume_text}\n\n"
        f"✅ Add a section for key achievements!\n"
    )
    return tips, new_resume_text

# ---------- Routes ----------

@app.route('/upload', methods=['POST'])
def upload_file():
    print ("✅ Received a request.")
    """
    Handle file upload and generate improved resume + feedback.
    """
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in request'}), 400

    file = request.files['file']
    job_role = request.form.get('jobRole', '').strip()

    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    if not job_role:
        return jsonify({'error': 'Job role is required'}), 400

    try:
        # Extract text and generate feedback
        resume_text = extract_text(file)
        if not resume_text.strip():
            return jsonify({'error': 'Failed to extract text from the PDF'}), 400

        tips, new_resume_text = generate_feedback(resume_text, job_role)

        # Save enhanced resume as a Word document
        output_path = 'improved_resume.docx'
        doc = Document()
        doc.add_heading('Improved Resume', 0)
        doc.add_paragraph(new_resume_text)
        doc.save(output_path)

        return jsonify({
            'tips': tips,
            'download_link': '/download'
        })

    except Exception as e:
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500

@app.route('/download', methods=['GET'])
def download_file():
    """
    Serve the improved resume file for download.
    """
    output_path = 'improved_resume.docx'
    if os.path.exists(output_path):
        return send_file(output_path, as_attachment=True)
    else:
        return jsonify({'error': 'No resume file available for download.'}), 404

@app.route('/enhance', methods=['POST'])
def enhance():
    data = request.get_json()
    bullet = data.get('bullet')
    if not bullet:
        return jsonify({'error': 'No bullet point provided.'}), 400
    
    # Enhance the bullet point
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a professional resume expert. Improve the following resume bullet point."},
            {"role": "user", "content": bullet}
        ]
    )
    improved = response['choices'][0]['message']['content']
    return jsonify({'original': bullet, 'improved': improved})
# ---------- Run the App ----------

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')

